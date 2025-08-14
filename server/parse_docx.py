import os, re, math
from typing import cast, Optional, TypedDict, Union
from docx.document import Document
from datetime import datetime

DEBUG = os.environ.get('DEBUG', 'false').lower() in ['true', '1']

month_names = {
    'Jan': 1, 'January': 1,
    'Feb': 2, 'February': 2,
    'Mar': 3, 'March': 3,
    'Apr': 4, 'April': 4,
    'May': 5,
    'Jun': 6, 'June': 6,
    'Jul': 7, 'July': 7,
    'Aug': 8, 'August': 8,
    'Sep': 9, 'September': 9,
    'Oct': 10, 'October': 10,
    'Nov': 11, 'November': 11,
    'Dec': 12, 'December': 12
}
# Regular expressions for date patterns
re_dd = fr'(?P<day>{ '|'.join([ f'0?{i}' for i in range(1, 10) ]) }|{ '|'.join([ str(i) for i in range(10, 32) ]) })'  # day
re_mm = fr'(?P<month>{ '|'.join([ f'0?{i}' for i in range(1, 13) ]) })'  # month
re_yy = r'(?P<year>\d{2,4})'  # year
re_months = r'(?P<month_name>Jan(uary)?|Feb(ruary)?|Mar(rch)?|Apr(il)?|May|June?|July?|Aug(ust)?|Sep(tember)?|Oct(ober)?|Nov(ember)?|Dec(ember)?)' # months
re_ordinal = r'(?:st|nd|rd|th)'  # ordinal suffixes
re_of = r'(?:,|\s+of\s+)'  # optional "of" word
# Possible patterns in order of "correctness"
date_patterns = [
    fr'\b{re_dd}[/-](?:{re_mm}|{re_months})[/-]{re_yy}\b',                      # DD/MM/YYYY or DD-MM-YYYY or DD/MONTH/YYYY DD-MONTH-YYYY
    fr'\b{re_yy}[/-](?:{re_mm}|{re_months})[/-]{re_dd}\b',                      # YYYY/MM/DD or YYYY-MM-DD or YYYY/MONTH/DD or YYYY-MONTH-DD
    fr'\b(?:{re_mm}|{re_months})[/-]{re_dd}[/-]{re_yy}\b',                      # MM/DD/YYYY or MM-DD-YYYY or MONTH/DD/YYYY MONTH-DD-YYYY
    fr'\b{re_yy}[/-]{re_dd}[/-](?:{re_mm}|{re_months})\b',                      # YYYY/DD/MM or YYYY-DD-MM or YYYY/DD/MONTH or YYYY-DD-MONTH
    fr'\b{re_dd}\s?{re_ordinal}?{re_of}?{re_months}(,|\s[a-z]*)?\s+{re_yy}\b',  # DD[...] Month[...] YYYY
    fr'\b{re_months}(?:\sthe)?\s+{re_dd}{re_of}?\s*{re_yy}\b',                  # Month [the] DD, YYYY
]

# Compile patterns for efficiency
compiled_patterns = [re.compile(pattern, re.IGNORECASE) for pattern in date_patterns]

# Types
class FoundDateMatch(TypedDict):
    found_date: str
    context: str

class FoundDateContext(TypedDict):
    type: str
    location: str
    text: str

FoundDate = Union[FoundDateMatch, FoundDateContext]

def _search_date_patterns(text: str):
    found_dates = []
    if not text.strip(): return found_dates

    found_positions = set()  # To avoid duplicate matches
    for pattern in compiled_patterns:
        # don't match parts of the text that were already matched by previous patterns
        matches = pattern.finditer(text)
        for match in matches:
            match_start = match.start()
            if match_start not in found_positions:
                found_dates.append((match_start, match))
                found_positions.add(match_start)

    # Return the matches in the order they appear in the text
    sorted_dates = [
        date for _start, date in sorted(found_dates, key=lambda x: x[0])
    ]
    return sorted_dates


# Date parsing logic
def _find_dates_in_text(text, data: FoundDateContext, context_limit = 20):
    """Return a list of found dates in the given text with context and additional data."""
    found_dates: list[FoundDate] = []
    if not text.strip(): return found_dates

    if DEBUG: print(f"Searching for dates in text ({data['type']}): {text[:context_limit]}...")
    century = math.floor(datetime.now().year / 100) * 100
    matches = _search_date_patterns(text)
    for match in matches:
        if DEBUG: print(f"Found date: {match.group(0)} in text: {data['text']}")
        try:
            match_groups = match.groupdict()
            found_date = datetime(
                year = int(match_groups['year']) if len(match_groups['year']) == 4 else int(match_groups['year']) + century,
                month = int(match_groups['month']) if 'month' in match_groups and match_groups['month'] is not None else month_names[match_groups['month_name']],
                day=int(match_groups['day']),
                hour=0, minute=0, second=0, microsecond=0
            )
            if DEBUG: print(f"Parsed date: {found_date}", match_groups)
        except ValueError as e:
            print(f"Invalid date found: {match.group(0)} - {e}")
            continue  # Skip invalid dates
        found_dates.append(cast(FoundDate, {
            'found_date': found_date.isoformat()+"Z",
            'found_text': match.group(0),
            'context': text[max(0, match.start() - context_limit // 2): min(len(text), match.end() + context_limit // 2)],
            **data
        }))
    return found_dates

# Document parsing logic
def find_dates_in_docx(doc: Document) -> list[FoundDate]:
    """Return a list of found dates in the given docx document."""
    found_dates = []

    # Look through section objects and their headers and footers 
    # https://python-docx.readthedocs.io/en/latest/api/section.html
    for section_index, section in enumerate(doc.sections):
        if section.different_first_page_header_footer:
            for header_index, header in enumerate(section.first_page_header.paragraphs):
                found_dates.extend(
                    _find_dates_in_text(header.text, data={
                        'type': 'header',
                        'location': f'section {section_index + 1}, first page header {header_index + 1}',
                        'text': header.text
                    })
                )
            for footer_index, footer in enumerate(section.first_page_footer.paragraphs):
                found_dates.extend(
                    _find_dates_in_text(footer.text, data={
                        'type': 'footer',
                        'location': f'section {section_index + 1}, first page footer {footer_index + 1}',
                        'text': footer.text
                    })
                )
        for header_index, header in enumerate(section.header.paragraphs):
            found_dates.extend(
                _find_dates_in_text(header.text, data={
                    'type': 'header',
                    'location': f'section {section_index + 1}, header {header_index + 1}',
                    'text': header.text
                })
            )
        for footer_index, footer in enumerate(section.footer.paragraphs):
            found_dates.extend(
                    _find_dates_in_text(footer.text, data={
                        'type': 'footer',
                        'location': f'section {section_index + 1}, footer {footer_index + 1}',
                        'text': footer.text
                    })
            )

    # Look through paragraphs and their text runs
    # https://python-docx.readthedocs.io/en/latest/api/text.html#paragraph-objects
    # https://python-docx.readthedocs.io/en/latest/api/text.html#docx.text.run.Run
    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        for run_index, run in enumerate(paragraph.runs):
            found_dates.extend(
                _find_dates_in_text(run.text, data={
                    'type': 'run',
                    'location': f'paragraph {paragraph_index + 1}, run {run_index + 1}',
                    'text': run.text
                })
            )

    # Look through tables, their cells and their paragraphs
    # https://python-docx.readthedocs.io/en/latest/api/table.html#docx.table._Cell.paragraphs
    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    for run_index, run in enumerate(paragraph.runs):
                        found_dates.extend(
                            _find_dates_in_text(run.text, data={
                                'type': 'run',
                                'location': f'table {table_index + 1}, row {row_index + 1}, cell {cell_index + 1}, paragraph {paragraph_index + 1}, run {run_index + 1}',
                                'text': run.text
                            })
                        )

    return found_dates
