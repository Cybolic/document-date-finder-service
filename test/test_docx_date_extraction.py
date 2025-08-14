import unittest
import re
from docx import Document
from datetime import datetime

from server.parse_docx import find_dates_in_docx

# Helper function to extract context around found text
def get_context(text: str, found_text: str, context_length: int = 10) -> str:
    return text[max(0, text.index(found_text) - context_length): min(len(text), text.index(found_text) + len(found_text) + context_length)]

class TestDocumentDateExtraction(unittest.TestCase):
    maxDiff = None

    def test_document_with_no_real_dates(self):
        document = Document()
        document.add_paragraph("A plain document containing no dates.")
        document.add_paragraph("Just some paragraphs and a few non-dates: 2025-08-135, 30/13/1996, Septober 6, 2000.")
        document.add_paragraph("Some more dates with bad numbers: 32-11-2022 / 04 / 00-04-00.")
        self.assertEqual(find_dates_in_docx(document), [])
    
    def test_document_with_one_date(self):
        date = '2023-10-01'
        document = Document()
        document.add_paragraph(f"A document with a single date on {date}.")

        found_dates = find_dates_in_docx(document)

        self.assertEqual(found_dates, [
            {
                'found_date': datetime(2023, 10, 1, 0, 0), 'found_text': date,
                'context': get_context(document.paragraphs[0].runs[0].text, date),
                'text': document.paragraphs[0].runs[0].text,
                'type': 'run', 'location': 'paragraph 1, run 1'
            }
        ])
    
    def test_document_with_multiple_dates(self):
        dates = [
            ('2023-10-01', datetime(2023, 10, 1, 0, 0)),
            ('2024-01-15', datetime(2024, 1, 15, 0, 0)),
            ('01-12-2025', datetime(2025, 12, 1, 0, 0))
        ]
        text = f"This document has multiple dates: {dates[0][0]}, {dates[1][0]}, and {dates[2][0]}."
        document = Document()
        document.add_paragraph(text)

        found_dates = find_dates_in_docx(document)

        self.assertEqual(len(found_dates), 3)
        self.assertEqual(found_dates, [
            {
                'found_date': dates[0][1], 'found_text': dates[0][0], 'text': text,
                'context': get_context(text, dates[0][0]),
                'type': 'run', 'location': 'paragraph 1, run 1',
            },
            {
                'found_date': dates[1][1], 'found_text': dates[1][0], 'text': text,
                'context': get_context(text, dates[1][0]),
                'type': 'run', 'location': 'paragraph 1, run 1',
            },
            {
                'found_date': dates[2][1], 'found_text': dates[2][0], 'text': text,
                'context': get_context(text, dates[2][0]),
                'type': 'run', 'location': 'paragraph 1, run 1',
            }
        ])

    def test_document_with_multiple_dates_in_paragraphs_and_tables(self):
        dates = [
            ('2023-10-01', datetime(2023, 10, 1, 0, 0)),
            ('5th of October, 2024', datetime(2024, 10, 5, 0, 0)),
            ('October 5 of 2024', datetime(2024, 10, 5, 0, 0)),
            ('31/12/2025', datetime(2025, 12, 31, 0, 0))
        ]
        text = f"This document has a date in the first paragraph: {dates[0][0]}."
        document = Document()
        document.add_paragraph(text)
        table = document.add_table(rows=3, cols=2)
        table.cell(0, 0).text = f"First cell with a date: {dates[1][0]}."
        table.cell(0, 1).text = "Second cell with no date."
        table.cell(1, 0).text = "Fourth cell with no date."
        table.cell(1, 1).text = f"Third cell with a date: {dates[2][0]}."
        table.cell(2, 0).text = f"Third cell with a date: {dates[3][0]}."
        table.cell(2, 1).text = "Fourth cell with no date."
        
        found_dates = find_dates_in_docx(document)
        
        self.assertEqual(found_dates, [
            {
                'found_date': dates[0][1], 'found_text': dates[0][0], 'text': text,
                'context': get_context(text, dates[0][0]),
                'type': 'run', 'location': 'paragraph 1, run 1',
            },
            {
                'found_date': dates[1][1], 'found_text': dates[1][0], 'text': document.tables[0].cell(0, 0).text,
                'context': get_context(document.tables[0].cell(0, 0).text, dates[1][0]),
                'type': 'run', 'location': 'table 1, row 1, cell 1, paragraph 1, run 1',
            },
            {
                'found_date': dates[2][1], 'found_text': dates[2][0], 'text': document.tables[0].cell(1, 1).text,
                'context': get_context(document.tables[0].cell(1, 1).text, dates[2][0]),
                'type': 'run', 'location': 'table 1, row 2, cell 2, paragraph 1, run 1',
            },
            {
                'found_date': dates[3][1], 'found_text': dates[3][0], 'text': document.tables[0].cell(2, 0).text,
                'context': get_context(document.tables[0].cell(2, 0).text, dates[3][0]),
                'type': 'run', 'location': 'table 1, row 3, cell 1, paragraph 1, run 1',
            }
        ])

    def test_document_with_dates_in_paragraphs_runs_inside_a_table(self):
        dates = [
            ('2023-10-01', datetime(2023, 10, 1, 0, 0)),
            ('2024-01-15', datetime(2024, 1, 15, 0, 0))
        ]
        document = Document()
        table = document.add_table(rows=1, cols=1)
        # note that adding a paragraph here means there's an empty one at the start of the cell
        paragraph = table.cell(0, 0).add_paragraph(f"This is a date inside a table: {dates[0][0]}.")
        paragraph.add_run(" And this is another date: 2024-01-15.")
        
        found_dates = find_dates_in_docx(document)
        
        self.assertEqual(len(found_dates), 2)
        self.assertEqual(found_dates, [
            {
                'found_date': dates[0][1], 'found_text': dates[0][0],
                'text': document.tables[0].cell(0, 0).paragraphs[1].runs[0].text,
                'context': get_context(document.tables[0].cell(0, 0).paragraphs[1].runs[0].text, dates[0][0]),
                'type': 'run', 'location': 'table 1, row 1, cell 1, paragraph 2, run 1',
            },
            {
                'found_date': dates[1][1], 'found_text': dates[1][0],
                'text': document.tables[0].cell(0, 0).paragraphs[1].runs[1].text,
                'context': get_context(document.tables[0].cell(0, 0).paragraphs[1].runs[1].text, dates[1][0]),
                'type': 'run', 'location': 'table 1, row 1, cell 1, paragraph 2, run 2',
            }
        ])

    def test_document_with_dates_in_headers_and_footers(self):
        dates = [
            ('2023-10-01', datetime(2023, 10, 1, 0, 0)),
            ('2024-01-15', datetime(2024, 1, 15, 0, 0))
        ]
        document = Document()
        section = document.sections[0]
        section.header.paragraphs[0].text = f"The original date is {dates[0][0]} in the morning."
        section.footer.paragraphs[0].text = f"Last updated: {dates[1][0]}."
        
        found_dates = find_dates_in_docx(document)
        
        self.assertEqual(len(found_dates), 2)
        self.assertEqual(found_dates, [
            {
                'found_date': dates[0][1], 'found_text': dates[0][0], 'text': section.header.paragraphs[0].text,
                'context': get_context(section.header.paragraphs[0].text, dates[0][0]),
                'type': 'header', 'location': 'section 1, header 1',
            },
            {
                'found_date': dates[1][1], 'found_text': dates[1][0], 'text': section.footer.paragraphs[0].text,
                'context': get_context(section.footer.paragraphs[0].text, dates[1][0]),
                'type': 'footer', 'location': 'section 1, footer 1',
            }
        ])

    def test_document_with_dates_in_first_page_header_and_footer(self):
        dates = [
            ('2023-10-01', datetime(2023, 10, 1, 0, 0)),
            ('2024-01-15', datetime(2024, 1, 15, 0, 0))
        ]
        document = Document()
        section = document.sections[0]
        section.different_first_page_header_footer = True
        section.first_page_header.paragraphs[0].text = f"First page header date: {dates[0][0]}."
        section.first_page_footer.paragraphs[0].text = f"First page footer date: {dates[1][0]}."
        
        found_dates = find_dates_in_docx(document)
        
        self.assertEqual(len(found_dates), 2)
        self.assertEqual(found_dates, [
            {
                'found_date': dates[0][1], 'found_text': dates[0][0], 'text': section.first_page_header.paragraphs[0].text,
                'context': get_context(section.first_page_header.paragraphs[0].text, dates[0][0]),
                'type': 'header', 'location': 'section 1, first page header 1',
            },
            {
                'found_date': dates[1][1], 'found_text': dates[1][0], 'text': section.first_page_footer.paragraphs[0].text,
                'context': get_context(section.first_page_footer.paragraphs[0].text, dates[1][0]),
                'type': 'footer', 'location': 'section 1, first page footer 1',
            }
        ])

    def test_document_with_various_date_formats(self):
        dates = [
          ('5-Feb-2019'         , '05-02-19 0:00'      , datetime(2019, 2, 5, 0, 0)),
          ('05/02/2019 00:00:00', '05/02/2019 00:00:00', datetime(2019, 2, 5, 0, 0)),
          ('5-Feb-2019'         , 'February 5, 2019'   , datetime(2019, 2, 5, 0, 0)),
          ('05-02-2019 00:00:00', 'February 5, 2019'   , datetime(2019, 2, 5, 0, 0)),
          ('05-06-19 0:00'      , 'June 5, 2019'       , datetime(2019, 6, 5, 0, 0)),
          ('05-06-19 0:00'      , '05/06/2019 00:00:00', datetime(2019, 6, 5, 0, 0)),
          ('05/16/2019 00:00:00', '05/16/2019 00:00:00', datetime(2019, 5, 16, 0, 0)),
          ('05/21/2019 00:00:00', '05/21/2019 00:00:00', datetime(2019, 5, 21, 0, 0)),
        ]
        document = Document()
        document.add_paragraph("Here are some dates in various formats:")
        for date1, date2, _date3 in dates:
            paragraph = document.add_paragraph(f"{date1},")
            paragraph.add_run(f" {date2}")

        found_dates = find_dates_in_docx(document)
        self.assertEqual(len(found_dates), 16)
        expected_dates = [
            ([
                {
                    'found_date': date[2], 'found_text': re.sub(r' [0:]*$', '', date[0]),
                    'context': f'{date[0]},',
                    'text': f"{date[0]},",
                    'type': 'run', 'location': f'paragraph {date_index + 2}, run 1'
                },
                {
                    'found_date': date[2], 'found_text': re.sub(r' [0:]*$', '', date[1]),
                    'context': f" {date[1]}",
                    'text': f" {date[1]}",
                    'type': 'run', 'location': f'paragraph {date_index + 2}, run 2'
                }
            ])
            for date_index, date in enumerate(dates)
        ]
        expected_dates_flattened = [ i for ii in expected_dates for i in ii ]

        self.assertEqual(found_dates, expected_dates_flattened)

if __name__ == '__main__':
    unittest.main()
